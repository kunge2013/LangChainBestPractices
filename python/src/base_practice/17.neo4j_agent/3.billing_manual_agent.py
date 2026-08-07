# [AGC:FILE] tool=Cc author=fangkun date=2026-08-07
"""
政企计费账务系统操作手册 - 多模态知识检索 Agent

基于 .docx 操作手册，提取文本和图片，用 qwen VL 模型生成图片描述，
存入 Neo4j 向量库，通过 deepagents 创建可自主检索的操作助手。

模块划分:
    Config         - 配置管理
    DocxParser     - 文档解析 (文本 + 图片提取)
    ImageDescriber - 图片描述生成
    VectorStore    - 向量库构建与检索
    BillingAgent   - Agent 创建与交互
"""

import base64
import io
import json
import logging
import os
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from dotenv import load_dotenv
from langchain_core.documents import Document as LangchainDocument
from langchain_core.messages import HumanMessage
from langchain_core.tools import BaseTool
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_neo4j.vectorstores.neo4j_vector import Neo4jVector
from langchain_openai import ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PIL import Image

from deepagents import create_deep_agent

load_dotenv()

# ======================== 日志配置 ========================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)


# ======================== 配置管理 ========================


@dataclass
class Config:
    """统一配置管理"""

    # 文档路径
    billing_manual_path: str = field(default_factory=lambda: os.path.join(
        os.path.dirname(__file__),
        "中国电信集团政企计费账务系统操作手册-客户经理分册 -修订-20260720.docx",
    ))
    images_output_dir: str = field(default_factory=lambda: os.path.join(
        os.path.dirname(__file__), "output", "images"
    ))

    # Neo4j
    vector_index_name: str = "billing_manual_vectors"
    neo4j_uri: str = field(default_factory=lambda: os.environ.get("NEO4J_URI", ""))
    neo4j_username: str = field(default_factory=lambda: os.environ.get("NEO4J_USERNAME", ""))
    neo4j_password: str = field(default_factory=lambda: os.environ.get("NEO4J_PASSWORD", ""))
    neo4j_database: str = field(default_factory=lambda: os.environ.get("NEO4J_DATABASE", "neo4j"))

    # 启动行为
    init_on_startup: bool = field(default_factory=lambda: os.environ.get("INIT_ON_STARTUP", "true").lower() == "true")

    # Embedding
    embedding_model: str = "shibing624/text2vec-base-chinese"
    embedding_device: str = "cpu"

    # 文本分块
    chunk_size: int = 1500
    chunk_overlap: int = 300

    # 检索
    search_k: int = 3
    image_base_url: str = field(default_factory=lambda: os.environ.get("IMAGE_BASE_URL", "http://localhost:2024"))

    # LLM (OpenAI 兼容协议)
    openai_model: str = field(default_factory=lambda: os.environ.get("OPENAI_MODEL", "qwen3.5-plus"))
    openai_base_url: str = field(default_factory=lambda: os.environ.get(
        "OPENAI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"
    ))
    openai_api_key: str = field(default_factory=lambda: os.environ.get("OPENAI_API_KEY", ""))
    openai_temperature: float = field(default_factory=lambda: float(os.environ.get("OPENAI_TEMPERATURE", "0.7")))
    openai_max_tokens: int = field(default_factory=lambda: int(os.environ.get("OPENAI_MAX_TOKENS", "2000")))
    describe_image_max_tokens: int = 102400
    describe_max_retries: int = 3

    def get_neo4j_params(self) -> dict:
        return {
            "url": self.neo4j_uri,
            "username": self.neo4j_username,
            "password": self.neo4j_password,
            "database": self.neo4j_database,
        }

    def get_llm_params(self) -> dict:
        return {
            "model": self.openai_model,
            "base_url": self.openai_base_url,
            "api_key": self.openai_api_key,
            "temperature": self.openai_temperature,
            "max_tokens": self.openai_max_tokens,
        }

    def image_to_html(self, image_path: str, alt: str = "") -> str:
        """将图片路径转为 HTML img 标签"""
        if not image_path:
            return ""
        filename = os.path.basename(image_path)
        return f'<img src="{self.image_base_url}/{filename}" alt="{alt}" />'


# ======================== 数据模型 ========================


@dataclass
class TextElement:
    type: str = "text"
    content: str = ""
    section: str = ""  # 所属章节
    is_heading: bool = False  # 是否为标题段落


@dataclass
class ImageElement:
    type: str = "image"
    path: str = ""
    nearest_paragraph_text: str = ""
    image_description: str = ""
    section: str = ""  # 所属章节


# ======================== 模块 1: 文档解析 ========================


class DocxParser:
    """解析 .docx 文件，提取文本段落和图片"""

    def __init__(self, config: Config):
        self.config = config

    def parse(self, docx_path: str) -> list[TextElement | ImageElement]:
        """解析文档，返回有序元素列表"""
        Path(self.config.images_output_dir).mkdir(parents=True, exist_ok=True)

        doc = Document(docx_path)
        elements: list[TextElement | ImageElement] = []
        image_counter = 0
        seen_rids = set()
        current_section = "前言"

        # 提取文本段落 + 段落内图片 (按文档顺序)
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = (para.style.name or "").lower()
            is_heading = "heading" in style_name or "title" in style_name

            if is_heading and text:
                current_section = text

            if text:
                elements.append(TextElement(content=text, section=current_section, is_heading=is_heading))

            # 从每个 run 中提取图片
            for run in para.runs:
                images = self._extract_images_from_run(run, doc, seen_rids)
                for image_bytes, image_ext in images:
                    image_path = os.path.join(
                        self.config.images_output_dir,
                        f"img_{image_counter}.{image_ext}",
                    )
                    try:
                        img = Image.open(io.BytesIO(image_bytes))
                        img.save(image_path, "PNG")
                        elements.append(ImageElement(
                            path=image_path,
                            nearest_paragraph_text=text if text else self._last_text(elements),
                            section=current_section,
                        ))
                        image_counter += 1
                        logger.info(f"提取图片 {image_counter}: {image_path} [{current_section}]")
                    except Exception as e:
                        logger.warning(f"图片保存失败: {e}")

        # 补充: 从 inline_shapes 提取
        for shape in doc.inline_shapes:
            image_bytes = self._extract_image_bytes(shape, seen_rids)
            if image_bytes is None:
                continue
            image_path = os.path.join(self.config.images_output_dir, f"img_{image_counter}.png")
            try:
                img = Image.open(io.BytesIO(image_bytes))
                img.save(image_path, "PNG")
                elements.append(ImageElement(
                    path=image_path,
                    nearest_paragraph_text=self._last_text(elements),
                    section=current_section,
                ))
                image_counter += 1
                logger.info(f"提取图片 {image_counter} (inline_shape): {image_path}")
            except Exception as e:
                logger.warning(f"图片保存失败: {e}")

        # 兜底: 遍历文档 part 的所有 image 关系
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            if rel.rId in seen_rids:
                continue
            try:
                blob = rel.target_part.blob
                ext = rel.target_part.content_type.split("/")[-1].split("+")[0] or "png"
                image_path = os.path.join(self.config.images_output_dir, f"img_{image_counter}.{ext}")
                img = Image.open(io.BytesIO(blob))
                img.save(image_path, "PNG")
                elements.append(ImageElement(
                    path=image_path,
                    nearest_paragraph_text=self._last_text(elements),
                    section=current_section,
                ))
                image_counter += 1
                logger.info(f"提取图片 {image_counter} (fallback): {image_path}")
            except Exception as e:
                logger.debug(f"兜底提取图片失败: {e}")

        text_count = len([e for e in elements if isinstance(e, TextElement)])
        image_count = len([e for e in elements if isinstance(e, ImageElement)])
        logger.info(f"文档解析完成: {text_count} 个文本段落, {image_count} 张图片")
        return elements

    @staticmethod
    def _last_text(elements: list) -> str:
        """返回最后一个文本元素的内容"""
        for elem in reversed(elements):
            if isinstance(elem, TextElement) and elem.content:
                return elem.content
        return ""

    @staticmethod
    def _extract_images_from_run(run, doc, seen_rids: set) -> list[tuple[bytes, str]]:
        """从 run 的 XML 中提取所有图片二进制数据"""
        images = []
        nsmap = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }

        blips = run._element.findall(".//a:blip", namespaces=nsmap)
        for blip in blips:
            rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rid or rid in seen_rids:
                continue

            try:
                image_part = doc.part.related_parts.get(rid)
                if image_part:
                    seen_rids.add(rid)
                    blob = image_part.blob
                    ext = image_part.content_type.split("/")[-1].split("+")[0] or "png"
                    images.append((blob, ext))
            except Exception as e:
                logger.debug(f"从 run 提取图片失败 (rid={rid}): {e}")

        return images

    @staticmethod
    def _extract_image_bytes(shape, seen_rids: set) -> bytes | None:
        """从 inline_shape 提取图片二进制数据"""
        try:
            nsmap = {
                "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            }
            r_ids = shape._element.findall(
                "pic:pic/pic:blipFill/a:blip/@r:embed", namespaces=nsmap
            )
            if not r_ids:
                return None
            r_id = r_ids[0]
            if r_id in seen_rids:
                return None
            seen_rids.add(r_id)
            image_part = shape.part.related_parts.get(r_id)
            return image_part.blob if image_part else None
        except Exception as e:
            logger.debug(f"inline_shape 提取图片失败: {e}")
            return None


# ======================== 模块 2: 图片描述生成 ========================


class ImageDescriber:
    """调用 VL 模型为图片生成中文描述"""

    PROMPT = (
        "请用中文详细描述这张图片的内容，包括图片中的文字、界面元素、操作流程、步骤等。"
        "描述要准确完整，以便后续用于知识检索。"
    )

    def __init__(self, config: Config):
        self.config = config
        self.model = ChatOpenAI(**config.get_llm_params())

    def describe_batch(
        self, images: list[ImageElement], max_retries: int | None = None
    ) -> list[ImageElement]:
        """批量描述图片，带重试"""
        max_retries = max_retries or self.config.describe_max_retries
        total = len(images)

        for i, img in enumerate(images, 1):
            if not img.path or not os.path.exists(img.path):
                logger.warning(f"图片不存在，跳过: {img.path}")
                continue

            for attempt in range(1, max_retries + 1):
                try:
                    logger.info(f"正在描述图片 {i}/{total} (尝试 {attempt}/{max_retries})")
                    img.image_description = self._describe_one(img.path)
                    logger.info(f"描述成功: {img.image_description[:50]}...")
                    break
                except Exception as e:
                    logger.warning(f"描述失败 (尝试 {attempt}/{max_retries}): {e}")
                    if attempt < max_retries:
                        time.sleep(2 ** attempt)
            else:
                img.image_description = "描述生成失败"

        logger.info(f"图片描述完成: {len([i for i in images if i.image_description])}/{total}")
        return images

    def _describe_one(self, image_path: str) -> str:
        """描述单张图片"""
        with open(image_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode("utf-8")

        message = HumanMessage(
            content=[
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_b64}"}},
                {"type": "text", "text": self.PROMPT},
            ]
        )

        response = self.model.invoke([message], max_tokens=self.config.describe_image_max_tokens)
        content = response.content
        return content.strip() if content else "图片描述生成失败"


# ======================== 模块 3: 向量库构建与检索 ========================


class VectorStoreBuilder:
    """构建 LangChain Documents 并写入 Neo4j 向量库"""

    def __init__(self, config: Config):
        self.config = config
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
        )
        self.embeddings = HuggingFaceEmbeddings(
            model_name=config.embedding_model,
            model_kwargs={"device": config.embedding_device},
            encode_kwargs={"normalize_embeddings": True},
        )

    def build_documents(self, elements: list[TextElement | ImageElement]) -> list[LangchainDocument]:
        """将解析元素转为 LangChain Documents，附加图片描述和章节信息"""
        docs = []
        chunk_index = 0

        for elem in elements:
            if not isinstance(elem, TextElement):
                continue

            nearby_images = self._find_nearby_images(elem, elements)
            metadata = {
                "source": "billing_manual",
                "section": elem.section,
                "chunk_index": chunk_index,
                "image_descriptions": json.dumps(
                    [img.image_description for img in nearby_images],
                    ensure_ascii=False,
                ),
                "image_paths": json.dumps(
                    [img.path for img in nearby_images],
                    ensure_ascii=False,
                ),
            }

            docs.append(LangchainDocument(page_content=elem.content, metadata=metadata))
            chunk_index += 1

        return docs

    @staticmethod
    def _find_nearby_images(text_elem: TextElement, elements: list) -> list[ImageElement]:
        """找到文本段落之后、下一个文本段落之前的所有图片"""
        images = []
        found = False
        for elem in elements:
            if elem is text_elem:
                found = True
                continue
            if found:
                if isinstance(elem, ImageElement):
                    images.append(elem)
                elif isinstance(elem, TextElement):
                    break
        return images

    def build(self, docs: list[LangchainDocument], force_recreate: bool = False) -> Neo4jVector:
        """创建或加载 Neo4j 向量库"""
        if not force_recreate:
            try:
                db = Neo4jVector.from_existing_index(
                    embedding=self.embeddings,
                    index_name=self.config.vector_index_name,
                    text_node_property="text",
                    embedding_node_property="embedding",
                    **self.config.get_neo4j_params(),
                )
                logger.info("成功加载已有向量库")
                return db
            except Exception as e:
                logger.warning(f"加载向量库失败: {e}")
                raise RuntimeError(
                    "向量库不存在。请先运行 pipeline.build(force_recreate=True) 初始化数据。"
                ) from e

        if not docs:
            raise ValueError("force_recreate=True 时必须提供文档列表")

        logger.info(f"正在创建向量库: {len(docs)} 个文档...")
        db = Neo4jVector.from_documents(
            documents=docs,
            embedding=self.embeddings,
            index_name=self.config.vector_index_name,
            text_node_property="text",
            embedding_node_property="embedding",
            **self.config.get_neo4j_params(),
        )
        logger.info("向量库创建完成")
        return db


class KnowledgeSearcher(BaseTool):
    """知识检索工具 — 支持相似度检索 + 按章节检索"""

    name: str = "search_knowledge"
    description: str = (
        "检索政企计费账务系统操作手册的知识库。返回相关文本内容、所属章节和关联图片的路径。"
        "参数 query 是检索关键词或问题。"
        "如果要查询某个章节的所有内容，使用另一个工具 get_section_images。"
    )

    db: Neo4jVector | None = None
    config: Config | None = None

    def __init__(self, db: Neo4jVector, config: Config):
        super().__init__(db=db, config=config)

    def _run(self, query: str) -> str:
        results = self.db.similarity_search_with_score(query, k=self.config.search_k)

        if not results:
            return "未找到相关结果，请尝试不同的关键词。"

        output_parts = []
        for i, (doc, score) in enumerate(results, 1):
            sec = doc.metadata.get("section", "")
            section = [f"--- 结果 {i} (相似度: {score:.4f}) [章节: {sec}] ---"]
            section.append(f"文本内容:\n{doc.page_content}")

            image_descs = json.loads(doc.metadata.get("image_descriptions", "[]"))
            image_paths = json.loads(doc.metadata.get("image_paths", "[]"))

            if image_descs:
                section.append("\n相关图片描述:")
                for j, desc in enumerate(image_descs):
                    section.append(f"  图{j + 1}: {desc}")

            if image_paths:
                section.append("\n相关图片:")
                for j, path in enumerate(image_paths):
                    desc = image_descs[j] if j < len(image_descs) else ""
                    section.append(f"  {self.config.image_to_html(path, alt=desc)}")

            output_parts.append("\n".join(section))

        return "\n\n".join(output_parts)


class GetSectionImages(BaseTool):
    """按章节检索所有图片"""

    name: str = "get_section_images"
    description: str = (
        "获取指定章节的所有文本内容和图片路径。"
        "参数 section_name 是章节名称，如'客户管理'、'账单管理'等。"
        "返回该章节的所有文本段落和关联图片。"
    )

    db: Neo4jVector | None = None
    config: Config | None = None

    def __init__(self, db: Neo4jVector, config: Config):
        super().__init__(db=db, config=config)

    def _run(self, section_name: str) -> str:
        # 通过 Neo4j 直接查询该章节的所有文档
        try:
            from neo4j import GraphDatabase

            driver = GraphDatabase.driver(
                self.config.neo4j_uri,
                auth=(self.config.neo4j_username, self.config.neo4j_password),
            )
            with driver.session(database=self.config.neo4j_database) as session:
                result = session.run(
                    """
                    MATCH (n)
                    WHERE n.source = 'billing_manual'
                    AND toLower(n.section) CONTAINS toLower($section)
                    RETURN n.text AS text, n.image_descriptions AS image_descs,
                           n.image_paths AS image_paths, n.section AS section
                    ORDER BY n.chunk_index
                    """,
                    section=section_name,
                )
                records = list(result)
            driver.close()
        except Exception as e:
            return f"查询章节失败: {e}"

        if not records:
            # 尝试列出所有章节名称
            try:
                from neo4j import GraphDatabase
                driver = GraphDatabase.driver(
                    self.config.neo4j_uri,
                    auth=(self.config.neo4j_username, self.config.neo4j_password),
                )
                with driver.session(database=self.config.neo4j_database) as session:
                    result = session.run(
                        """
                        MATCH (n) WHERE n.source = 'billing_manual' AND n.section IS NOT NULL
                        RETURN DISTINCT n.section AS section ORDER BY n.chunk_index
                        """
                    )
                    sections = [r["section"] for r in result]
                driver.close()
                return f"未找到章节 '{section_name}'。可用章节:\n" + "\n".join(f"- {s}" for s in sections)
            except Exception:
                return f"未找到章节 '{section_name}'。"

        output = [f"## 章节: {records[0].get('section', '未知')}"]
        all_images = set()

        for rec in records:
            text = rec.get("text", "")
            if text:
                output.append(text[:500])

            img_paths = json.loads(rec.get("image_paths", "[]") or "[]")
            all_images.update(img_paths)

        if all_images:
            output.append(f"\n### 图片 ({len(all_images)} 张)")
            for path in sorted(all_images):
                output.append(self.config.image_to_html(path, alt=os.path.basename(path)))

        return "\n\n".join(output)


# ======================== 模块 4: Agent 创建 ========================


class BillingAgent:
    """政企计费账务系统操作助手"""

    SYSTEM_PROMPT = """你是一个政企计费账务系统操作助手，帮助用户理解和完成系统操作。

## 可用工具

1. **search_knowledge** — 根据用户问题关键词检索相关内容（向量相似度匹配）
2. **get_section_images** — 根据章节名称获取该章节全部内容和图片

## 使用策略

- 用户问**具体问题**（如"如何创建新客户"）→ 用 `search_knowledge`
- 用户要**整个章节**（如"给我看客户管理章节的图片"）→ 用 `get_section_images`
- 如果用户提到"章节"、"第X章"、"全部图片" → 优先用 `get_section_images`

## 输出规范

### 操作流程问题
- **文字步骤说明**
- **Mermaid 流程图**：用 ```mermaid 代码块
  ```mermaid
  flowchart TD
      A[步骤1] --> B[步骤2]
      B --> C[步骤3]
  ```
- **相关图片路径**

### 界面/布局问题
- 重点返回图片路径 + 文字描述

### 概念/定义问题
- 直接文字说明，无需流程图

## 图片格式
图片以 HTML img 标签形式返回，如 `<img src="http://localhost:2024/img_0.png" alt="描述" />`。
回答时请直接输出这些 img 标签，让 Markdown 渲染器展示图片。
请用中文回答所有问题。"""

    def __init__(self, config: Config, db: Neo4jVector):
        self.config = config
        self.db = db
        self.searcher = KnowledgeSearcher(db, config)
        self.section_tool = GetSectionImages(db, config)
        self.agent = None

    def create(self):
        """创建 deep_agent 实例"""
        model = ChatOpenAI(**self.config.get_llm_params())
        self.agent = create_deep_agent(
            model=model,
            tools=[self.searcher, self.section_tool],
            system_prompt=self.SYSTEM_PROMPT,
        )
        return self.agent

    def invoke(self, query: str) -> str:
        """调用 agent 处理用户查询"""
        if self.agent is None:
            self.create()

        result = self.agent.invoke({"messages": [{"role": "user", "content": query}]})

        if isinstance(result, dict) and "messages" in result:
            for msg in result["messages"]:
                if hasattr(msg, "content") and msg.content:
                    return msg.content
        return str(result)


# ======================== 主流程 ========================


class BillingManualPipeline:
    """完整 pipeline: 解析 → 描述 → 分块 → 建库 → 创建 agent"""

    def __init__(self, config: Config | None = None):
        self.config = config or Config()

    def build(self, force_recreate: bool = False) -> tuple[Neo4jVector, BillingAgent]:
        # 阶段 1: 解析文档
        self._log_stage(1, "解析 .docx 文档")
        parser = DocxParser(self.config)
        elements = parser.parse(self.config.billing_manual_path)

        # 阶段 2: 生成图片描述
        self._log_stage(2, "生成图片描述")
        describer = ImageDescriber(self.config)
        images = [e for e in elements if isinstance(e, ImageElement)]
        # 只保留三张图
        images = images[:3]
        describer.describe_batch(images)

        # 阶段 3: 构建向量库
        self._log_stage(3, "构建向量库")
        store_builder = VectorStoreBuilder(self.config)
        docs = store_builder.build_documents(elements)
        db = store_builder.build(docs, force_recreate)

        # 阶段 4: 创建 Agent
        self._log_stage(4, "创建 Agent")
        agent = BillingAgent(self.config, db)
        agent.create()

        return db, agent

    @staticmethod
    def _log_stage(stage: int, name: str):
        sep = "=" * 50
        logger.info(sep)
        logger.info(f"阶段 {stage}: {name}")
        logger.info(sep)


# ======================== LangGraph API 入口 ========================


def build_agent() -> Any:
    """
    LangGraph API 所需的 agent factory 函数。
    根据 config.init_on_startup 决定是否初始化向量库。
    """
    config = Config()

    if config.init_on_startup:
        # 启动时执行完整 pipeline（解析 → 描述 → 建库）
        pipeline = BillingManualPipeline(config)
        _, billing_agent = pipeline.build(force_recreate=False)
        return billing_agent.agent
    else:
        # 仅连接已有向量库
        store_builder = VectorStoreBuilder(config)
        db = store_builder.build([], force_recreate=False)
        billing_agent = BillingAgent(config, db)
        billing_agent.create()
        return billing_agent.agent


# 模块级 agent 变量，供 langgraph.json 引用
agent = build_agent()


# ======================== CLI 入口 ========================
#
#
# def main():
#     """
#     CLI 入口 — 仅用于初始化数据:
#         python 3.billing_manual_agent.py init      # 首次运行，解析+描述+建库
#         python 3.billing_manual_agent.py rebuild   # 强制重建向量库
#     """
#     import sys
#
#     # if len(sys.argv) < 2:
#     #     logger.info("用法: python 3.billing_manual_agent.py init [--rebuild]")
#     #     return
#
#     # command = sys.argv[1]
#     command = "init"
#     config = Config()
#
#     if command == "init":
#         force_recreate = len(sys.argv) > 2 and sys.argv[2] == "rebuild"
#         force_recreate = True
#         pipeline = BillingManualPipeline(config)
#         _, billing_agent = pipeline.build(force_recreate=force_recreate)
#         logger.info(f"数据准备完成! Agent 类型: {type(billing_agent.agent)}")
#     elif command == "rebuild":
#         pipeline = BillingManualPipeline(config)
#         _, billing_agent = pipeline.build(force_recreate=True)
#         logger.info(f"向量库重建完成! Agent 类型: {type(billing_agent.agent)}")
#     else:
#         logger.info(f"未知命令: {command}")
#         logger.info("用法: python 3.billing_manual_agent.py init [--rebuild]")


# if __name__ == "__main__":
#     main()
